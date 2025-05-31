import streamlit as st
from docx import Document
import io
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional
from io import BytesIO
import os


class DocumentProcessor:
    """Responsible for processing and filling a docx document with data."""

    @staticmethod
    def preencher_contrato_docx(
        dados_contrato: Dict[str, Any], caminho_template: str
    ) -> Optional[BytesIO]:
        """
        Fills a .docx template with contract data and returns the document in memory.

        Args:
            dados_contrato: Dictionary with keys matching placeholders in the template.
            caminho_template: Path to the template .docx file.

        Returns:
            BytesIO object containing the filled document or None if an error occurs.
        """
        try:
            documento = Document(caminho_template)

            # Process paragraphs
            for p in documento.paragraphs:
                DocumentProcessor._replace_placeholders_in_paragraph(p, dados_contrato)

            # Process tables
            for table in documento.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            DocumentProcessor._replace_placeholders_in_paragraph(
                                p, dados_contrato
                            )

            # Save to memory buffer
            bio = io.BytesIO()
            documento.save(bio)
            bio.seek(0)
            return bio

        except FileNotFoundError:
            st.error(
                f"Erro: O arquivo de template não foi encontrado em '{caminho_template}'."
            )
            return None
        except Exception as e:
            st.error(f"Ocorreu um erro inesperado ao gerar o documento: {e}")
            return None

    @staticmethod
    def _replace_placeholders_in_paragraph(paragraph, dados_contrato):
        """Replace placeholders in a paragraph with values from the data dictionary."""
        inline = paragraph.runs
        for i in range(len(inline)):
            texto_combinado = inline[i].text
            for chave, valor in dados_contrato.items():
                placeholder_docx = f"{{{{{chave}}}}}"  # Format {{key}}
                if placeholder_docx in texto_combinado:
                    texto_combinado = texto_combinado.replace(
                        placeholder_docx, str(valor)
                    )
            inline[i].text = texto_combinado


class FormSection(ABC):
    """Abstract base class for form sections."""

    @abstractmethod
    def render(self, col, dados: Dict[str, Any]) -> None:
        """Render the form section and collect data."""
        pass


class VendedorSection(FormSection):
    """Form section for seller data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Dados do Vendedor 1", expanded=True):
                dados["vendedor_1_nome"] = st.text_input("Nome do Vendedor 1", "JOSÉ")
                dados["vendedor_1_nacionalidade"] = st.text_input(
                    "Nacionalidade do Vendedor 1", "brasileiro"
                )
                dados["vendedor_1_estado_civil"] = st.text_input(
                    "Estado Civil do Vendedor 1", "convivente"
                )
                dados["vendedor_1_profissao"] = st.text_input(
                    "Profissão do Vendedor 1", "motorista"
                )
                dados["vendedor_1_rg"] = st.text_input(
                    "RG do Vendedor 1", "X.XXX.XXX-X"
                )
                dados["vendedor_1_cpf"] = st.text_input(
                    "CPF do Vendedor 1", "XXX.XXX.XXX-XX"
                )
                dados["vendedor_1_endereco"] = st.text_input(
                    "Endereço do Vendedor 1",
                    "Rua Da Esperança, 1. Maringá/PR",
                )

            with st.expander("Dados da Vendedora 2", expanded=True):
                dados["vendedor_2_nome"] = st.text_input(
                    "Nome da Vendedora 2", "SOLANGE"
                )
                dados["vendedor_2_nacionalidade"] = st.text_input(
                    "Nacionalidade da Vendedora 2", "brasileira"
                )
                dados["vendedor_2_estado_civil"] = st.text_input(
                    "Estado Civil da Vendedora 2", "convivente"
                )
                dados["vendedor_2_profissao"] = st.text_input(
                    "Profissão da Vendedora 2", "do lar"
                )
                dados["vendedor_2_rg"] = st.text_input(
                    "RG da Vendedora 2", "X.XXX.XXX-X"
                )
                dados["vendedor_2_cpf"] = st.text_input(
                    "CPF da Vendedora 2", "XXX.XXX.XXX-XX"
                )
                dados["vendedor_2_endereco"] = st.text_input(
                    "Endereço da Vendedora 2",
                    "Rua Do Amor, 1. Maringá/PR",
                )


class CompradorSection(FormSection):
    """Form section for buyer data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Dados do Comprador", expanded=True):
                dados["comprador_nome"] = st.text_input("Nome do Comprador", "JOÃO")
                dados["comprador_nacionalidade"] = st.text_input(
                    "Nacionalidade do Comprador", "brasileiro"
                )
                dados["comprador_profissao"] = st.text_input(
                    "Profissão do Comprador", "auxiliar de escritório"
                )
                dados["comprador_rg"] = st.text_input("RG do Comprador", "X.XXX.XXX-X")
                dados["comprador_cpf"] = st.text_input(
                    "CPF do Comprador", "XXX.XXX.XXX-XX"
                )
                dados["comprador_endereco"] = st.text_input(
                    "Endereço do Comprador",
                    "Rua Do Sucesso, 1. Maringá/PR",
                )

            with st.expander("Dados da Procuradora do Comprador"):
                dados["procurador_nome"] = st.text_input("Nome da Procuradora", "MARIA")
                dados["procurador_nacionalidade"] = st.text_input(
                    "Nacionalidade da Procuradora", "brasileira"
                )
                dados["procurador_estado_civil"] = st.text_input(
                    "Estado Civil da Procuradora", "casada"
                )
                dados["procurador_profissao"] = st.text_input(
                    "Profissão da Procuradora", "do lar"
                )
                dados["procurador_rg"] = st.text_input(
                    "RG da Procuradora", "X.XXX.XXX-X"
                )
                dados["procurador_cpf"] = st.text_input(
                    "CPF da Procuradora", "XXX.XXX.XXX-XX"
                )
                dados["procurador_endereco"] = st.text_input(
                    "Endereço da Procuradora",
                    "Rua Do Sucesso, 1. Maringá/PR",
                )


class ImovelSection(FormSection):
    """Form section for property data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Detalhes do Imóvel", expanded=True):
                dados["imovel_descricao"] = st.text_input(
                    "Descrição do Imóvel",
                    'casa "2" (dois) do CONDOMINIO RESIDENCIAL CLARO E PIMENTEL',
                )
                dados["imovel_area_privativa"] = st.text_input(
                    "Área Privativa (m²)", "99,95"
                )
                dados["imovel_area_terreno"] = st.text_input(
                    "Fração Ideal do Terreno (m²)", "150,15"
                )
                dados["imovel_endereco"] = st.text_input(
                    "Endereço do Imóvel",
                    "Rua Do Sucesso, 2. Maringá/PR",
                )
                dados["imovel_matricula"] = st.text_input(
                    "Matrícula do Imóvel", "12.345"
                )
                dados["imovel_serventia_registral"] = st.text_input(
                    "Serventia Registral",
                    "3ª. Serventia Registral da Comarca de Maringá/PR",
                )


class FinanciamentoSection(FormSection):
    """Form section for financing data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Detalhes do Financiamento", expanded=True):
                dados["imovel_banco_financiamento"] = st.text_input(
                    "Banco do Financiamento", "Caixa Econômica Federal"
                )
                dados["imovel_prazo_financiamento_meses"] = st.text_input(
                    "Prazo do Financiamento (meses)", "420"
                )
                dados["imovel_prazo_financiamento_extenso"] = st.text_input(
                    "Prazo por Extenso", "quatrocentos e vinte meses"
                )
                dados["imovel_parcelas_vencidas"] = st.text_input(
                    "Nº de Parcelas Vencidas", "21"
                )
                dados["imovel_parcelas_vencidas_extenso"] = st.text_input(
                    "Nº Parcelas Vencidas por Extenso", "vinte e uma"
                )
                dados["imovel_mes_base_vencimento"] = st.text_input(
                    "Mês Base das Parcelas Vencidas", "maio de 2016"
                )


class PagamentoSection(FormSection):
    """Form section for payment data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Valores e Forma de Pagamento", expanded=True):
                dados["pagamento_valor_total_direitos"] = st.text_input(
                    "Valor Total dos Direitos (R$)", "40.000,00"
                )
                dados["pagamento_valor_total_direitos_extenso"] = st.text_input(
                    "Valor Total por Extenso", "quarenta mil reais"
                )
                st.markdown("---")
                dados["pagamento_parcela_1_valor"] = st.text_input(
                    "Valor Parcela 1 (R$)", "15.000,00"
                )
                dados["pagamento_parcela_1_extenso"] = st.text_input(
                    "Parcela 1 por Extenso", "quinze mil reais"
                )
                dados["pagamento_parcela_1_data"] = st.text_input(
                    "Data Parcela 1", "06/05/2000"
                )
                dados["pagamento_parcela_1_banco"] = st.text_input(
                    "Banco Parcela 1", "CEF"
                )
                dados["pagamento_parcela_1_agencia"] = st.text_input("Agência", "XXXX")
                dados["pagamento_parcela_1_operacao"] = st.text_input("Operação", "XXX")
                dados["pagamento_parcela_1_conta"] = st.text_input(
                    "Conta/Poupança", "XXXX-X"
                )
                dados["pagamento_parcela_1_favorecido"] = st.text_input(
                    "Favorecido Parcela 1", "José"
                )
                st.markdown("---")
                dados["pagamento_parcela_2_valor"] = st.text_input(
                    "Valor Parcela 2 - Comissão (R$)", "3.000,00"
                )
                dados["pagamento_parcela_2_extenso"] = st.text_input(
                    "Parcela 2 por Extenso", "três mil reais"
                )
                dados["pagamento_parcela_2_favorecido"] = st.text_input(
                    "Favorecido Parcela 2", "Richard"
                )
                st.markdown("---")
                dados["pagamento_parcela_3_valor"] = st.text_input(
                    "Valor Parcela 3 - Vencimento (R$)", "2.000,00"
                )
                dados["pagamento_parcela_3_extenso"] = st.text_input(
                    "Parcela 3 por Extenso", "dois mil reais"
                )
                dados["pagamento_parcela_3_vencimento"] = st.text_input(
                    "Mês de Vencimento Parcela 3", "maio"
                )
                st.markdown("---")
                dados["pagamento_parcela_4_valor"] = st.text_input(
                    "Valor Parcela 4 - Final (R$)", "20.000,00"
                )
                dados["pagamento_parcela_4_extenso"] = st.text_input(
                    "Parcela 4 por Extenso", "vinte mil reais"
                )
                dados["pagamento_parcela_4_data"] = st.text_input(
                    "Data Parcela 4", "31/12/2000"
                )


class ResponsabilidadesSection(FormSection):
    """Form section for responsibilities data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Responsabilidades (IPTU e Financiamento)", expanded=True):
                dados["iptu_responsabilidade_vendedor_ate"] = st.text_input(
                    "Vendedor paga IPTU até", "dezembro de 2000"
                )
                dados["iptu_data_desocupacao"] = st.text_input(
                    "Data de Desocupação", "dezembro de 2000"
                )
                dados["iptu_responsabilidade_comprador_desde"] = st.text_input(
                    "Comprador paga IPTU a partir de", "janeiro de 2001"
                )
                dados["financiamento_responsabilidade_comprador_desde"] = st.text_input(
                    "Comprador paga financiamento a partir de", "junho de 2000"
                )
                dados["financiamento_procurador_cartao"] = st.text_input(
                    "Procurador(a) com cartão do débito", "RICHARD"
                )
                dados["financiamento_agencia_debito"] = st.text_input(
                    "Agência para Débito", "XXX"
                )
                dados["financiamento_operacao_debito"] = st.text_input(
                    "Operação para Débito", "XXX"
                )
                dados["financiamento_conta_debito"] = st.text_input(
                    "Conta para Débito", "XXXX-X"
                )
                dados["financiamento_procurador_venda"] = st.text_input(
                    "Procurador(a) para venda futura", "RICHARD"
                )


class TestemunhasSection(FormSection):
    """Form section for witnesses data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Testemunhas", expanded=True):
                dados["testemunha_1_nome"] = st.text_input(
                    "Nome da Testemunha 1", "VICTOR"
                )
                dados["testemunha_1_cpf"] = st.text_input(
                    "CPF da Testemunha 1", "XXX.XXX.XXX-XX"
                )
                st.markdown("---")
                dados["testemunha_2_nome"] = st.text_input(
                    "Nome da Testemunha 2", "SOFIA"
                )
                dados["testemunha_2_cpf"] = st.text_input(
                    "CPF da Testemunha 2", "XXX.XXX.XXX-XX"
                )


class DadosGeraisSection(FormSection):
    """Form section for general contract data."""

    def render(self, col, dados: Dict[str, Any]) -> None:
        with col:
            with st.expander("Dados Gerais do Contrato", expanded=True):
                dados["numero_registro"] = st.text_input("Número de Registro", "XXXXX")
                dados["cidade_contrato"] = st.text_input(
                    "Cidade de Assinatura", "Maringá"
                )
                dados["uf_contrato"] = st.text_input("UF de Assinatura", "PR")
                dados["estado_contrato"] = st.text_input("Estado por Extenso", "Paraná")
                dados["data_assinatura"] = st.text_input(
                    "Data de Assinatura", "06 de maio de 2000"
                )


class ContratoApp:
    """Main application class for contract generation."""

    def __init__(self):
        self.template_path = "contrato_particular_de_compromisso_de_compra_e_venda.docx"
        self.dados = {}

    def setup_page(self):
        """Configure the Streamlit page."""
        st.set_page_config(
            layout="wide", page_title="Gerador de Contrato de Compra e Venda"
        )
        st.title("📄 Gerador de Contrato de Compra e Venda")
        st.markdown(
            "Preencha os campos abaixo para gerar o contrato em formato `.docx`."
        )

    def render_form(self):
        """Render the form and handle submission."""
        with st.form(key="contrato_form"):
            # First column sections
            col1, col2 = st.columns(2)

            with col1:
                st.header("Partes Envolvidas")
                VendedorSection().render(col1, self.dados)
                CompradorSection().render(col1, self.dados)

            with col2:
                st.header("Dados do Imóvel e Contrato")
                ImovelSection().render(col2, self.dados)
                FinanciamentoSection().render(col2, self.dados)
                PagamentoSection().render(col2, self.dados)

            # Additional sections
            st.header("Cláusulas Adicionais e Assinaturas")
            col3, col4 = st.columns(2)

            ResponsabilidadesSection().render(col3, self.dados)
            TestemunhasSection().render(col4, self.dados)
            DadosGeraisSection().render(col4, self.dados)

            # Submit button
            submitted = st.form_submit_button("Gerar Contrato .DOCX")

        return submitted

    def generate_contract(self):
        """Generate contract document and handle download."""

        print(os.getcwd())

        st.error(f"O caminho atual do script é: {os.getcwd()}")
        documento_gerado_buffer = DocumentProcessor.preencher_contrato_docx(
            self.dados, self.template_path
        )

        if documento_gerado_buffer:
            st.success("✅ Contrato gerado com sucesso!")

            # Create filename dynamically
            nome_arquivo = f"Contrato_{self.dados.get('comprador_nome', 'comprador').replace(' ', '_')}.docx"

            st.download_button(
                label="Clique aqui para baixar o Contrato (.docx)",
                data=documento_gerado_buffer,
                file_name=nome_arquivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    def run(self):
        """Run the application."""
        self.setup_page()
        submitted = self.render_form()

        if submitted:
            self.generate_contract()


if __name__ == "__main__":
    app = ContratoApp()
    app.run()
